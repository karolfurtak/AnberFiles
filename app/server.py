#!/usr/bin/env python3
"""File server dla /mnt/data/sprawozdania/ — port 8765, HTTP basic auth.

Konfiguracja przez env (systemd EnvironmentFile=/etc/sprawozdania-server.env):
    SERVER_PORT (default 8765)
    SERVER_USER (default anbernic)
    SERVER_PASS — WYMAGANE
    SERVER_HOST (default 0.0.0.0 = LAN, można 127.0.0.1 = tylko SSH tunnel)

Listing katalogu = sortowalna tabela (Nazwa, Rozmiar, Modyfikacja, Utworzono).
Kliknięcie nagłówka kolumny sortuje (rozmiar i daty sortowane numerycznie).
"""
import os
import re
import html as _html
import base64
from datetime import datetime
from pathlib import Path
from urllib.parse import quote
import asyncio
import hashlib
from aiohttp import web

try:
    import markdown as _markdown      # renderowany podgląd .md (opcjonalny)
except Exception:
    _markdown = None


def _battery_html() -> str:
    """Poziom baterii konsoli (PMIC axp2202) do linii informacyjnej —
    aktualizuje się razem z auto-odświeżaniem listingu."""
    try:
        base = Path('/sys/class/power_supply/axp2202-battery')
        cap = int((base / 'capacity').read_text().strip())
        st = (base / 'status').read_text().strip()
    except Exception:
        return ''
    ikona = '⚡' if st in ('Charging', 'Full') else '🔋'
    col = '#1d7a36' if cap > 40 else ('#9a7b00' if cap > 15 else '#c00')
    return (f' · <span style="color:{col};font-weight:600">'
            f'{ikona} {cap}%</span>')


def _natkey(s: str):
    """Klucz sortowania naturalnego: 'plik_10' PO 'plik_9' (liczby jako liczby)."""
    return [int(p) if p.isdigit() else p.lower() for p in re.split(r'(\d+)', s)]

ROOT  = Path('/mnt/data/sprawozdania')
PORT  = int(os.environ.get('SERVER_PORT', '8765'))
HOST  = os.environ.get('SERVER_HOST', '0.0.0.0')
USER  = os.environ.get('SERVER_USER', 'anbernic')
PASSW = os.environ.get('SERVER_PASS', '')

ICONS = {'pdf': '📕', 'html': '🌐', 'md': '📝', 'jpg': '🖼', 'jpeg': '🖼',
         'png': '🖼', 'xlsx': '📊', 'xls': '📊', 'csv': '📊', 'txt': '📄',
         'docx': '📘', 'doc': '📘', 'svg': '🖼', 'json': '🔧'}

IMG_EXT = {'.png', '.jpg', '.jpeg', '.gif', '.webp', '.bmp', '.svg'}
AUDIO_EXT = {'.mp3', '.flac', '.wav', '.ogg', '.m4a', '.opus'}
AUDIO_MIME = {'.mp3': 'audio/mpeg', '.flac': 'audio/flac', '.wav': 'audio/wav',
              '.ogg': 'audio/ogg', '.m4a': 'audio/mp4', '.opus': 'audio/opus'}

AUDIO_STYLE = (
    'body{margin:0;background:#14161c;color:#cfd6df;font-family:system-ui,sans-serif;'
    'display:flex;flex-direction:column;min-height:100vh}'
    '.bar{display:flex;gap:.6em;align-items:center;padding:.55em 1em;'
    'background:#1d2027;flex-wrap:wrap}'
    '.bar a,.bar button{color:#7ab7ff;background:none;border:1px solid #343a45;'
    'border-radius:5px;padding:.25em .7em;text-decoration:none;cursor:pointer;'
    'font:inherit;font-size:.95em}'
    '.bar a:hover,.bar button:hover{background:#272b34}'
    '.bar button.on{background:#1a5fb4;border-color:#1a5fb4;color:#fff}'
    '.wrap{flex:1;display:flex;flex-direction:column;align-items:center;'
    'justify-content:center;gap:1.2em;padding:1em}'
    '.tname{font-size:1.15em;color:#eee;word-break:break-all;text-align:center;'
    'max-width:90%}'
    'audio{width:min(680px,92vw)}'
    '.spd{display:flex;gap:.4em;align-items:center;color:#8a93a0;font-size:.9em}'
)


def render_audio_page(target: Path) -> str:
    """Odtwarzacz audio (?view=1): natywny <audio>, poprzedni/następny
    w katalogu, regulacja tempa (lektor!), autoodtwarzanie."""
    sibs = sorted((x.name for x in target.parent.iterdir()
                   if x.is_file() and x.suffix.lower() in AUDIO_EXT
                   and not x.name.startswith('.')), key=_natkey)
    idx = sibs.index(target.name) if target.name in sibs else 0
    prv = quote(sibs[idx - 1]) + '?view=1' if idx > 0 else None
    nxt = quote(sibs[idx + 1]) + '?view=1' if idx < len(sibs) - 1 else None
    a_prev = (f'<a href="{prv}" id="prev">← poprzedni</a>' if prv
              else '<a style="opacity:.3;pointer-events:none">← poprzedni</a>')
    a_next = (f'<a href="{nxt}" id="next">następny →</a>' if nxt
              else '<a style="opacity:.3;pointer-events:none">następny →</a>')
    q = quote(target.name)
    spd_btns = ''.join(
        f'<button class="sp" data-s="{s}">{s}×</button>'
        for s in ('0.75', '1', '1.25', '1.5', '2'))
    return (
        '<!doctype html><meta charset=utf-8>'
        '<meta name="viewport" content="width=device-width,initial-scale=1">'
        f'<title>♪ {target.name}</title><style>{AUDIO_STYLE}</style>'
        f'<div class="bar"><a href="./">📁 folder</a>{a_prev}{a_next}'
        f'<span style="color:#8a93a0">{idx + 1} / {len(sibs)}</span>'
        f'<a href="{q}?dl=1">⬇ pobierz</a></div>'
        f'<div class="wrap"><div class="tname">🎵 {target.name}</div>'
        f'<audio id="au" controls autoplay src="{q}"></audio>'
        f'<div class="spd">tempo: {spd_btns}</div></div>'
        '<script>(function(){'
        'const au=document.getElementById("au");'
        'const sp=[...document.querySelectorAll(".sp")];'
        'function setr(r){au.playbackRate=parseFloat(r);'
        'sp.forEach(b=>b.classList.toggle("on",b.dataset.s===r));'
        'try{localStorage.setItem("audioRate",r)}catch(e){}}'
        'sp.forEach(b=>b.onclick=()=>setr(b.dataset.s));'
        'setr(localStorage.getItem("audioRate")||"1");'
        # koniec utworu → automatycznie następny (playlista po katalogu)
        'au.onended=()=>{const n=document.getElementById("next");if(n)location=n.href};'
        'document.addEventListener("keydown",e=>{'
        'if(e.key==="ArrowLeft"){const a=document.getElementById("prev");if(a)location=a.href}'
        'if(e.key==="ArrowRight"){const a=document.getElementById("next");if(a)location=a.href}'
        'if(e.key===" "){e.preventDefault();au.paused?au.play():au.pause()}});'
        '})();</script>')

DOCX_CACHE = Path('/mnt/data/.cache/docx-preview')
_LO_LOCK = asyncio.Lock()   # jedna konwersja naraz (A53)


async def docx_to_pdf(target: Path) -> Path | None:
    """Podglad .docx: konwersja LibreOffice -> PDF, cache per (sciezka, mtime).
    Pierwsze otwarcie ~8-12 s, kolejne natychmiast."""
    DOCX_CACHE.mkdir(parents=True, exist_ok=True)
    key = hashlib.sha1(str(target).encode()).hexdigest()[:16]
    cached = DOCX_CACHE / f'{key}_{int(target.stat().st_mtime)}.pdf'
    if cached.exists():
        return cached
    # sprzatnij stare wersje tego pliku
    for old_f in DOCX_CACHE.glob(f'{key}_*.pdf'):
        try:
            old_f.unlink()
        except Exception:
            pass
    async with _LO_LOCK:
        if cached.exists():
            return cached
        proc = await asyncio.create_subprocess_exec(
            'soffice', '--headless',
            '-env:UserInstallation=file:///tmp/lo_preview_profile',
            '--convert-to', 'pdf', '--outdir', str(DOCX_CACHE), str(target),
            stdout=asyncio.subprocess.DEVNULL, stderr=asyncio.subprocess.DEVNULL)
        try:
            await asyncio.wait_for(proc.wait(), timeout=90)
        except asyncio.TimeoutError:
            proc.kill()
            return None
    produced = DOCX_CACHE / (target.stem + '.pdf')
    if produced.exists():
        produced.rename(cached)
        return cached
    return None


MD_STYLE = (
    'body{margin:0;background:#f7f7f9;color:#222;font-family:system-ui,sans-serif}'
    '.bar{position:sticky;top:0;display:flex;gap:.6em;align-items:center;'
    'padding:.5em 1em;background:#26292f;color:#ddd;flex-wrap:wrap;z-index:10}'
    '.bar a,.bar button{color:#7ab7ff;background:none;border:1px solid #3a3f47;'
    'border-radius:5px;padding:.25em .7em;text-decoration:none;cursor:pointer;'
    'font:inherit;font-size:.95em}'
    '.bar a:hover,.bar button:hover{background:#32363d}'
    '.bar .name{color:#eee;font-weight:600;word-break:break-all}'
    '.bar button.on{background:#1a5fb4;border-color:#1a5fb4;color:#fff}'
    '#rendered{max-width:900px;margin:0 auto;padding:1.5em 2em;background:#fff;'
    'min-height:92vh;box-shadow:0 0 12px rgba(0,0,0,.07)}'
    '#rendered img{max-width:100%}'
    '#rendered table{border-collapse:collapse;margin:.8em 0}'
    '#rendered th,#rendered td{border:1px solid #999;padding:.35em .6em}'
    '#rendered code{background:#eef1f6;padding:0 .25em;border-radius:3px}'
    '#rendered pre{background:#1b1d21;color:#e8e8e8;padding:1em;border-radius:6px;'
    'overflow-x:auto}'
    '#rendered pre code{background:none}'
    '#rendered blockquote{border-left:4px solid #bcd;margin-left:0;'
    'padding-left:1em;color:#555}'
    '#raw{display:none;max-width:1100px;margin:0 auto;padding:1em}'
    '#raw pre{background:#1b1d21;color:#d8e0ea;padding:1.2em;border-radius:8px;'
    'overflow-x:auto;font-size:.9em;line-height:1.45;white-space:pre-wrap;'
    "font-family:'Cascadia Mono',Consolas,monospace}"
)


def _fix_md_imgs(body: str, md_dir: Path) -> str:
    """Obrazki w .md bywaja zapisane gola nazwa, a fizycznie leza w raw/
    obok (md w processed/). Znajdz plik i przepisz src na sciezke wzgledna
    dzialajaca z URL-a strony podgladu."""
    def repl(m):
        src = m.group(2)
        if src.startswith(('http://', 'https://', 'data:', '/')):
            return m.group(0)
        name = Path(src).name
        cands = [md_dir / src, md_dir / name,
                 md_dir.parent / 'raw' / name, md_dir / 'raw' / name,
                 md_dir.parent / 'exports' / name,
                 md_dir.parent / 'processed' / name]
        for c in cands:
            try:
                cr = c.resolve()
            except Exception:
                continue
            if cr.exists() and (ROOT in cr.parents):
                rel = os.path.relpath(cr, md_dir).replace(os.sep, '/')
                return m.group(1) + quote(rel) + m.group(3)
        return m.group(0)
    return re.sub(r'(<img[^>]*?src=")([^"]+)(")', repl, body)


def render_md_page(target: Path) -> str:
    """Podgląd .md: zakładki Render (markdown+MathJax) / Kod (surowe źródło).
    Względne ścieżki obrazków działają — strona żyje w katalogu pliku."""
    src = target.read_text(encoding='utf-8', errors='replace')
    if _markdown is not None:
        body = _markdown.markdown(
            src, extensions=['tables', 'fenced_code', 'nl2br', 'sane_lists'])
        body = _fix_md_imgs(body, target.parent)
    else:
        body = '<p><i>(brak biblioteki python-markdown — dostępny tylko Kod)</i></p>'
    raw = _html.escape(src)
    q = quote(target.name)
    # nawigacja po plikach .md w tym katalogu (jak w viewerze zdjec)
    sibs = sorted((x.name for x in target.parent.iterdir()
                   if x.is_file() and x.suffix.lower() == '.md'
                   and not x.name.startswith('.')), key=_natkey)
    idx = sibs.index(target.name) if target.name in sibs else 0
    prv = quote(sibs[idx - 1]) + '?view=1' if idx > 0 else None
    nxt = quote(sibs[idx + 1]) + '?view=1' if idx < len(sibs) - 1 else None
    a_prev = (f'<a href="{prv}" id="prev">← poprzedni</a>' if prv
              else '<a style="opacity:.3;pointer-events:none">← poprzedni</a>')
    a_next = (f'<a href="{nxt}" id="next">następny →</a>' if nxt
              else '<a style="opacity:.3;pointer-events:none">następny →</a>')
    return (
        '<!doctype html><meta charset=utf-8>'
        '<meta name="viewport" content="width=device-width,initial-scale=1">'
        f'<title>{target.name}</title><style>{MD_STYLE}</style>'
        '<script>window.MathJax={tex:{inlineMath:[["$","$"]],'
        'displayMath:[["$$","$$"]]}};</script>'
        '<script async src="https://cdn.jsdelivr.net/npm/mathjax@3/es5/'
        'tex-mml-chtml.js"></script>'
        f'<div class="bar"><a href="./">📁 folder</a>{a_prev}{a_next}'
        f'<button id="bren" class="on">Render</button>'
        f'<button id="braw">Kod</button>'
        f'<span class="name">{target.name}</span>'
        f'<span style="color:#888">{idx + 1} / {len(sibs)}</span>'
        f'<a href="{q}?dl=1">⬇ pobierz</a>'
        f'<a href="#" id="prn" data-n="{q}">🖨 drukuj</a></div>'
        f'<div id="rendered">{body}</div>'
        f'<div id="raw"><pre>{raw}</pre></div>'
        '<script>(function(){'
        'const r=document.getElementById("rendered"),'
        'w=document.getElementById("raw"),'
        'br=document.getElementById("bren"),bw=document.getElementById("braw");'
        'function show(ren){r.style.display=ren?"block":"none";'
        'w.style.display=ren?"none":"block";'
        'br.classList.toggle("on",ren);bw.classList.toggle("on",!ren);}'
        'br.onclick=()=>show(true);bw.onclick=()=>show(false);'
        # druk na Canon G3070 (md→DOCX→PDF→lp na konsoli, ~1-2 min)
        'document.getElementById("prn").onclick=async e=>{'
        'e.preventDefault();const a=e.target;'
        'if(!confirm("Wydrukować na Canon G3070?\\n'
        '(md→DOCX→PDF jak przy sprawozdaniach; ok. 1–2 min; '
        'drukarka musi być w sieci domowej)"))return;'
        'a.textContent="⏳...";'
        'try{const r=await fetch(a.dataset.n+"?print=1",{method:"POST"});'
        'a.textContent=r.status===202?"🖨 wysłano":"🖨 błąd";}'
        'catch(err){a.textContent="🖨 błąd";}'
        'setTimeout(()=>a.textContent="🖨 drukuj",4000);};'
        # AUTO-ODSWIEZANIE: poll mtime co 3 s; przy zmianie pobierz strone,
        # podmien tresc obu widokow i przelicz wzory MathJax. Zakladka
        # i pozycja przewiniecia zostaja.
        'let _mt=0;'
        'async function chk(){try{'
        'const j=await(await fetch(location.pathname+"?mt=1",'
        '{cache:"no-store"})).json();'
        'if(_mt&&j.mt!==_mt){'
        'const doc=new DOMParser().parseFromString('
        'await(await fetch(location.pathname+"?view=1",'
        '{cache:"no-store"})).text(),"text/html");'
        'const nr=doc.getElementById("rendered"),nw=doc.getElementById("raw");'
        'if(nr)r.innerHTML=nr.innerHTML;'
        'if(nw)w.innerHTML=nw.innerHTML;'
        'if(window.MathJax&&MathJax.typesetPromise)MathJax.typesetPromise([r]);}'
        '_mt=j.mt;}catch(e){}}'
        'setInterval(chk,3000);chk();'
        'document.addEventListener("keydown",e=>{'
        'if(e.key==="ArrowLeft"){const a=document.getElementById("prev");'
        'if(a)location=a.href}'
        'if(e.key==="ArrowRight"){const a=document.getElementById("next");'
        'if(a)location=a.href}});'
        '})();</script>')

STYLE = (
    'body{font-family:system-ui,sans-serif;max-width:1000px;margin:1.2em auto;'
    'padding:0 1em;background:#f7f7f9;color:#222}'
    'h2{color:#333;font-size:1.1em;word-break:break-all}'
    'table{width:100%;border-collapse:collapse;background:#fff;border-radius:6px;'
    'overflow:hidden;box-shadow:0 1px 3px rgba(0,0,0,.08)}'
    'th,td{text-align:left;padding:.5em .7em;border-bottom:1px solid #eee;'
    'font-size:.92em;white-space:nowrap}'
    'th{background:#eef2f7;cursor:pointer;user-select:none}'
    'th:hover{background:#e0e8f0}th::after{content:" \\2195";color:#aaa;font-size:.8em}'
    'td:first-child,th:first-child{white-space:normal;word-break:break-all}'
    'tr:hover td{background:#f3f7fc}'
    'a{text-decoration:none;color:#0066cc}a:hover{text-decoration:underline}'
    '.dir{font-weight:600}'
    'td:nth-child(2),th:nth-child(2){text-align:right;color:#555}'
    '.muted{color:#999;font-size:.85em;margin:.6em 0}'
    'h2 a{color:#0066cc}h2 .sep{color:#bbb;font-weight:400}'
    '.dl{margin-right:.45em;text-decoration:none;opacity:.55}'
    '.dl:hover{opacity:1;text-decoration:none}'
    '.del:hover{filter:drop-shadow(0 0 2px #c00)}'
    '.crumb{position:relative;display:inline-block}'
    '.crumb>.dd{display:none;position:absolute;top:100%;left:0;background:#fff;'
    'border:1px solid #c5cdd8;border-radius:6px;box-shadow:0 4px 14px rgba(0,0,0,.18);'
    'z-index:60;max-height:320px;overflow-y:auto;min-width:200px;padding:.25em 0}'
    '.crumb:hover>.dd{display:block}'
    '.crumb>.dd a{display:block;padding:.3em .9em;white-space:nowrap;font-size:.85em}'
    '.crumb>.dd a:hover{background:#eef2f7}'
    '.crumb>.dd a.cur{font-weight:700;color:#1a5fb4}'
    # TRYB MOBILNY (telefon): prostszy interfejs — bez usuwania/zmiany
    # nazwy/kopiowania nazwy; większe pola dotyku; mniej kolumn
    '@media (max-width:700px){'
    'body{margin:.6em auto;padding:0 .5em}'
    'a.del,a.ren,a.cpy{display:none}'
    'th:nth-child(4),td:nth-child(4){display:none}'   # kolumna Utworzono
    'th,td{padding:.85em .55em;font-size:1.02em}'
    '.dl{font-size:1.35em;padding:.15em .25em;margin-right:.4em}'
    'h2{font-size:1.05em}'
    '.crumb>.dd a{padding:.7em 1em;font-size:1em}'
    '}'
    '#dropov{position:fixed;inset:0;display:flex;align-items:center;'
    'justify-content:center;background:rgba(15,50,110,.88);color:#fff;'
    'font-size:1.5em;z-index:50;visibility:hidden;pointer-events:none;'
    'text-align:center;padding:1em}'
    '#dropov.on{visibility:visible}'
)

# Akcje plikowe: 🗑 usuń (→.kosz), ✎ zmień nazwę (prompt → POST ?rename=),
# ⧉ kopiuj nazwę do schowka.
DEL_JS = (
    '<script>'
    # modal wyboru formatu — RADIOBUTTONY per format + Generuj/Anuluj
    'function pickFmt(name){return new Promise(res=>{'
    'const ov=document.createElement("div");'
    'ov.style.cssText="position:fixed;inset:0;background:rgba(0,0,0,.5);'
    'display:flex;align-items:center;justify-content:center;z-index:99";'
    'const opts=[["","⚙ wg ustawień lektora (#e-lektor-ustawienia)"],'
    '["mp3","MP3 — 96 kbps (najmniejszy plik)"],'
    '["flac","FLAC — bezstratny zapis źródła 96 kbps"],'
    '["wav","WAV — nieskompresowany PCM"]];'
    'const dopts=[["","⚙ wg ustawień lektora"],'
    '["tak","TAK — model wizyjny opisze każdą ilustrację (dłużej)"],'
    '["nie","NIE — czytaj tylko podpisy rysunków"]];'
    'function mkradios(arr,nm){let h="";'
    'for(const [v,l] of arr){h+=\'<label style="display:flex;'
    'align-items:center;gap:.55em;padding:.45em .3em;cursor:pointer;'
    'border-radius:6px" onmouseover="this.style.background=\\\'#eef2f7\\\'" '
    'onmouseout="this.style.background=\\\'\\\'">'
    '<input type="radio" name="\'+nm+\'" value="\'+v+\'"\'+(v===""?" checked":"")'
    '+\' style="accent-color:#1a5fb4;width:1.05em;height:1.05em">\'+l'
    '+\'</label>\';}return h;}'
    'const radios=mkradios(opts,"lfmt"),dradios=mkradios(dopts,"lopis");'
    'ov.innerHTML=\'<div style="background:#fff;border-radius:10px;'
    'padding:1.1em 1.4em;max-width:92vw;min-width:300px;'
    'box-shadow:0 8px 30px rgba(0,0,0,.35)">'
    '<div style="font-weight:600;margin-bottom:.35em;word-break:break-all">'
    '🔊 Lektor: \'+name+\'</div>'
    '<div style="color:#556;font-size:.9em;margin-bottom:.7em">'
    'Format nagrania (generacja dłuższych dokumentów może potrwać '
    'kilkanaście minut):</div>\'+radios+'
    '\'<div style="color:#556;font-size:.9em;margin:.8em 0 .3em;'
    'border-top:1px solid #e4e8ee;padding-top:.7em">'
    '🖼 Opisy ilustracji (model wizyjny):</div>\'+dradios+'
    '\'<div style="display:flex;gap:.6em;margin-top:1em;'
    'justify-content:flex-end">'
    '<button data-a="x">Anuluj</button>'
    '<button data-a="ok" style="background:#1a5fb4;color:#fff;'
    'border-color:#1a5fb4">Generuj</button></div></div>\';'
    'ov.querySelectorAll("button").forEach(b=>{'
    'b.style.cssText+=";padding:.5em 1.2em;border:1px solid #b8c0cc;'
    'border-radius:6px;cursor:pointer;font-size:1em"'
    '+(b.dataset.a==="x"?";background:#f2f5f9":"");'
    'b.onclick=()=>{const v=ov.querySelector("input[name=lfmt]:checked");'
    'const d=ov.querySelector("input[name=lopis]:checked");'
    'ov.remove();res(b.dataset.a==="ok"'
    '?{fmt:(v?v.value:""),opisy:(d?d.value:"")}:null);};});'
    'ov.onclick=e=>{if(e.target===ov){ov.remove();res(null);}};'
    'document.body.appendChild(ov);});}'
    'document.addEventListener("click",async e=>{'
    'const a=e.target.closest("a.del,a.ren,a.cpy,a.lek");if(!a)return;'
    'e.preventDefault();'
    'const n=decodeURIComponent(a.dataset.n);'
    'if(a.classList.contains("lek")){'
    'const fm=await pickFmt(n);'
    'if(fm===null)return;'
    'a.textContent="⏳";'
    'try{'
    'let u=a.dataset.n+"?lektor=1"+(fm.fmt?"&fmt="+fm.fmt:"")'
    '+(fm.opisy?"&opisy="+fm.opisy:"");'
    'let r=await fetch(u,{method:"POST"});'
    'let j=await r.json().catch(()=>({}));'
    'if(j.status==="busy"){'
    'a.textContent="🔊";'
    'if(!confirm("Lektor jest teraz zajęty — czyta inny dokument'
    '(np. na polecenie z Discorda)."+(j.pending?"\\nW kolejce czeka: "'
    '+j.pending+" plik(ów).":"")+"\\n\\nDopisać ten plik do kolejki? '
    'Audio wygeneruje się automatycznie, gdy przyjdzie jego kolej."))return;'
    'a.textContent="⏳";'
    'r=await fetch(u+"&queue=1",{method:"POST"});'
    'j=await r.json().catch(()=>({}));}'
    'if(r.status===202){'
    'a.title=(j.status==="queued"?"W kolejce (poz. "+j.position+"): "'
    ':"Generuję: ")+(j.out||"");'
    'setTimeout(()=>a.textContent="🔊",2500);}'
    'else{alert("Lektor: "+(j.status||r.status));a.textContent="🔊";}'
    '}catch(err){alert("Błąd lektora");a.textContent="🔊";}return;}'
    'if(a.classList.contains("cpy")){'
    'try{await navigator.clipboard.writeText(n);'
    'a.textContent="✓";setTimeout(()=>a.textContent="⧉",900);}'
    'catch(err){prompt("Skopiuj nazwę:",n);}return;}'
    'if(a.classList.contains("ren")){'
    'const nn=prompt("Nowa nazwa:",n);'
    'if(!nn||nn===n)return;'
    'try{const r=await fetch(a.dataset.n+"?rename="+encodeURIComponent(nn),'
    '{method:"POST"});'
    'if(r.ok){location.reload();}'
    'else{alert("Błąd zmiany nazwy: "+await r.text());}'
    '}catch(err){alert("Błąd zmiany nazwy");}return;}'
    'if(!confirm("Usunąć \\""+n+"\\"?\\n(plik trafi do kosza .kosz)"))return;'
    'try{const r=await fetch(a.dataset.n,{method:"DELETE"});'
    'if(r.ok){a.closest("tr").remove();}'
    'else{alert("Błąd usuwania: HTTP "+r.status);}'
    '}catch(err){alert("Błąd usuwania");}'
    '});</script>'
)

# Drag & drop upload — upuszczenie plików na listing wgrywa je do bieżącego
# katalogu (POST multipart); tabela odświeży się sama (auto-refresh).
DROP_JS = (
    '<script>(function(){'
    'const ov=document.createElement("div");ov.id="dropov";'
    'ov.textContent="⬆ Upuść pliki, aby wgrać do tego katalogu";'
    'document.body.appendChild(ov);let d=0;'
    'window.addEventListener("dragenter",e=>{e.preventDefault();d++;ov.classList.add("on");});'
    'window.addEventListener("dragleave",e=>{e.preventDefault();d--;'
    'if(d<=0){d=0;ov.classList.remove("on");}});'
    'window.addEventListener("dragover",e=>e.preventDefault());'
    'window.addEventListener("drop",async e=>{e.preventDefault();d=0;'
    'const fs=[...e.dataTransfer.files];'
    'if(!fs.length){ov.classList.remove("on");return;}'
    'ov.textContent="⬆ Wgrywam "+fs.length+" plik(ów)…";'
    'const fd=new FormData();fs.forEach(f=>fd.append("file",f,f.name));'
    'try{const r=await fetch(location.pathname,{method:"POST",body:fd});'
    'const j=await r.json();'
    'ov.textContent=r.ok?("✓ Wgrano: "+(j.saved||[]).join(", ")):"✗ Błąd wgrywania";'
    '}catch(err){ov.textContent="✗ Błąd wgrywania";}'
    'setTimeout(()=>{ov.classList.remove("on");'
    'ov.textContent="⬆ Upuść pliki, aby wgrać do tego katalogu";},2500);'
    '});})();</script>'
)

VIEWER_STYLE = (
    'body{margin:0;background:#1b1d21;color:#ddd;font-family:system-ui,sans-serif;'
    'display:flex;flex-direction:column;min-height:100vh}'
    '.bar{display:flex;align-items:center;gap:1em;padding:.55em 1em;background:#26292f;'
    'font-size:.95em;flex-wrap:wrap}'
    '.bar a{color:#7ab7ff;text-decoration:none;padding:.25em .7em;border:1px solid #3a3f47;'
    'border-radius:5px;white-space:nowrap}'
    '.bar a:hover{background:#32363d}'
    '.bar .name{color:#eee;font-weight:600;word-break:break-all}'
    '.bar .cnt{color:#888}'
    '.stage{flex:1;display:flex;align-items:center;justify-content:center;'
    'padding:.8em;overflow:hidden;cursor:grab}'
    '.stage.drag{cursor:grabbing}'
    '.stage img{max-width:100%;max-height:calc(100vh - 5em);'
    'box-shadow:0 2px 14px rgba(0,0,0,.5);transform-origin:center center;'
    'user-select:none;-webkit-user-drag:none}'
    '.nav-off{opacity:.3;pointer-events:none}'
)

# Zoom: kółko / CTRL+kółko wokół kursora, drag = pan, dwuklik = fit/100%,
# klawisze +/-/0; % powiększenia widoczny na pasku (#zl).
VIEWER_JS = (
    '<script>(function(){'
    'const st=document.querySelector(".stage"),im=st.querySelector("img"),'
    'zl=document.getElementById("zl");'
    'let s=1,tx=0,ty=0;'
    'function ap(){im.style.transform=`translate(${tx}px,${ty}px) scale(${s})`;'
    'zl.textContent=Math.round(s*100)+"%";}'
    'function clamp(){s=Math.min(Math.max(s,0.2),20);}'
    'st.addEventListener("wheel",e=>{e.preventDefault();'
    'const r=im.getBoundingClientRect(),'
    'cx=e.clientX-(r.left+r.width/2),cy=e.clientY-(r.top+r.height/2),'
    'k=e.deltaY<0?1.2:1/1.2,o=s;s*=k;clamp();const f=s/o;'
    'tx=tx*f - cx*(f-1);ty=ty*f - cy*(f-1);ap();},{passive:false});'
    'let dr=null;'
    'st.addEventListener("mousedown",e=>{dr={x:e.clientX-tx,y:e.clientY-ty};'
    'st.classList.add("drag");e.preventDefault();});'
    'window.addEventListener("mousemove",e=>{if(!dr)return;'
    'tx=e.clientX-dr.x;ty=e.clientY-dr.y;ap();});'
    'window.addEventListener("mouseup",()=>{dr=null;st.classList.remove("drag");});'
    'function fit(){s=1;tx=ty=0;ap();}'
    'function nat(){s=im.naturalWidth&&im.clientWidth?im.naturalWidth/im.clientWidth:1;'
    'clamp();tx=ty=0;ap();}'
    # dwuklik: powiększone/przesunięte -> dopasuj do okna; dopasowane -> 1:1
    'st.addEventListener("dblclick",()=>{if(s!==1||tx||ty){fit();}else{nat();}});'
    'const fb=document.getElementById("fitb");if(fb)fb.onclick=e=>{e.preventDefault();fit();};'
    'const nb=document.getElementById("natb");if(nb)nb.onclick=e=>{e.preventDefault();nat();};'
    'document.addEventListener("keydown",e=>{'
    'if(e.key==="+"||e.key==="="){s*=1.2;clamp();ap();}'
    'if(e.key==="-"){s/=1.2;clamp();ap();}'
    'if(e.key==="0"){fit();}'
    'if(e.key==="1"){nat();}});'
    'ap();})();</script>'
)

# Sortowanie zapamiętywane w localStorage (przeżywa odświeżenie i nawigację);
# wiersz ".." (class="up") zawsze przypięty na górze.
SORT_JS = (
    '<script>(function(){'
    'const ths=[...document.querySelectorAll("th")];'
    'const tb=document.querySelector("table").tBodies[0];'
    # sortowanie NATURALNE tekstów: liczby w nazwach porównywane numerycznie
    'function nat(s){return s.split(/(\\d+)/).map('
    'p=>/^\\d+$/.test(p)?p.padStart(14,"0"):p.toLowerCase()).join("\\u0001");}'
    'function srt(i,asc){'
    'const up=tb.querySelector("tr.up");'
    '[...tb.rows].filter(r=>!r.classList.contains("up")).sort((a,b)=>{'
    'const dx=a.cells[i].dataset.sort,dy=b.cells[i].dataset.sort;'
    'let x=dx!==undefined?parseFloat(dx):nat(a.cells[i].textContent.trim());'
    'let y=dy!==undefined?parseFloat(dy):nat(b.cells[i].textContent.trim());'
    'return (x>y?1:x<y?-1:0)*(asc?1:-1);'
    '}).forEach(r=>tb.appendChild(r));'
    'if(up)tb.prepend(up);}'
    'ths.forEach((th,i)=>{th.onclick=()=>{'
    'const asc=th._a=!th._a;srt(i,asc);'
    'try{localStorage.setItem("dirSort",JSON.stringify({i:i,asc:asc}))}catch(e){}'
    '};});'
    'window._applySaved=function(){try{const s=JSON.parse(localStorage.getItem("dirSort"));'
    'if(s&&ths[s.i]!==undefined){ths[s.i]._a=s.asc;srt(s.i,s.asc);}}catch(e){}};'
    'window._applySaved();'
    # AUTO-ODŚWIEŻANIE: co 4 s pobierz w tle tę samą stronę, porównaj <tbody>
    # serwerowe z poprzednim pobraniem; przy zmianie podmień tabelę i licznik,
    # po czym przywróć zapamiętane sortowanie. Nowe pliki (np. świeży DOCX)
    # pojawiają się bez ręcznego odświeżania.
    '(function(){let sig="";'
    'async function tick(){try{'
    'const r=await fetch(location.pathname+location.search,{cache:"no-store"});'
    'if(!r.ok)return;'
    'const doc=new DOMParser().parseFromString(await r.text(),"text/html");'
    'const nb=doc.querySelector("tbody");if(!nb)return;'
    'const ns=nb.innerHTML;'
    'if(sig&&ns!==sig){'
    'document.querySelector("tbody").innerHTML=ns;'
    'const m=doc.querySelector("p.muted"),lm=document.querySelector("p.muted");'
    'if(m&&lm)lm.innerHTML=m.innerHTML;'
    'window._applySaved();}'
    'sig=ns;}catch(e){}}'
    'setInterval(tick,4000);})();'
    '})();</script>'
)


@web.middleware
async def auth(request, handler):
    # Pusty SERVER_PASS = brak autoryzacji (open access — tylko dla prywatnych sieci!)
    if not PASSW:
        return await handler(request)
    h = request.headers.get('Authorization', '')
    if not h.startswith('Basic '):
        return web.Response(status=401, headers={'WWW-Authenticate': 'Basic realm="Anbernic"'})
    try:
        u, p = base64.b64decode(h[6:]).decode().split(':', 1)
    except Exception:
        return web.Response(status=401, headers={'WWW-Authenticate': 'Basic realm="Anbernic"'})
    if u != USER or p != PASSW:
        return web.Response(status=401, text='Niepoprawne dane',
                            headers={'WWW-Authenticate': 'Basic realm="Anbernic"'})
    return await handler(request)


def _fmt_size(s):
    if s >= 1 << 30:
        return f'{s/(1<<30):.1f} GB'
    if s >= 1 << 20:
        return f'{s>>20} MB'
    if s >= 1024:
        return f'{s>>10} KB'
    return f'{s} B'


def _fmt_time(ts):
    try:
        return datetime.fromtimestamp(ts).strftime('%Y-%m-%d %H:%M')
    except Exception:
        return '—'


async def serve(request):
    # widok/stan kolejki lektora (dostępny z dowolnej ścieżki)
    if 'lektorq' in request.query:
        return web.Response(text=LEKTORQ_PAGE, content_type='text/html')
    if 'lektorqj' in request.query:
        return web.json_response(_lektor_queue_json())

    raw = request.match_info.get('path', '').strip('/')
    try:
        target = (ROOT / raw).resolve()
    except Exception:
        return web.Response(status=400)
    # path traversal guard
    if ROOT not in target.parents and target != ROOT:
        return web.Response(status=403)
    if not target.exists():
        return web.Response(status=404, text=f'Not found: {raw}')

    if target.is_dir():
        items = sorted(target.iterdir(), key=lambda p: (not p.is_dir(), _natkey(p.name)))
        rel = target.relative_to(ROOT)
        rel_url = f'/{rel}/' if str(rel) != '.' else '/'

        # breadcrumb — każdy segment klikalny; najechanie rozwija listę
        # folderów-RODZEŃSTWA z danego poziomu (skok w inną gałąź drzewa)
        crumbs = ['<a href="/" class="dir">📁 root</a>']
        acc = ''
        if str(rel) != '.':
            parent = ROOT
            for seg in str(rel).replace('\\', '/').split('/'):
                parent_url = quote(acc)          # ścieżka rodzica ('' dla 1. poziomu)
                acc += '/' + seg
                try:
                    sibs = sorted((d.name for d in parent.iterdir()
                                   if d.is_dir() and not d.name.startswith('.')),
                                  key=_natkey)
                except Exception:
                    sibs = []
                dd_items = []
                for s in sibs:
                    cls = ' class="cur"' if s == seg else ''
                    dd_items.append(f'<a href="{parent_url}/{quote(s)}/"{cls}>📁 {s}</a>')
                dd = (f'<div class="dd">{"".join(dd_items)}</div>') if dd_items else ''
                crumbs.append(f'<span class="crumb">'
                              f'<a href="{quote(acc)}/">{seg}</a>{dd}</span>')
                parent = parent / seg
        breadcrumb = ' <span class="sep">/</span> '.join(crumbs) + ' <span class="sep">/</span>'

        rows = []
        n = 0
        if target != ROOT:
            rows.append('<tr class="up"><td><a href="../" class="dir">📁 ..</a></td>'
                        '<td data-sort="-2">—</td><td data-sort="0"></td>'
                        '<td data-sort="0"></td></tr>')
        for item in items:
            if item.name.endswith(('.meta.json', '.resume.json')) or item.name.startswith('.'):
                continue
            try:
                st = item.stat()
                bt = getattr(st, 'st_birthtime', st.st_ctime)   # crtime jeśli dostępny, inaczej ctime
                mt_s, bt_s = _fmt_time(st.st_mtime), _fmt_time(bt)
                if item.is_dir():
                    rows.append(
                        f'<tr><td><a href="{item.name}/" class="dir">📁 {item.name}/</a></td>'
                        f'<td data-sort="-1">—</td>'
                        f'<td data-sort="{st.st_mtime:.0f}">{mt_s}</td>'
                        f'<td data-sort="{bt:.0f}">{bt_s}</td></tr>')
                else:
                    s = st.st_size
                    icon = ICONS.get(item.suffix.lower().lstrip('.'), '📄')
                    q = quote(item.name)
                    # zdjęcia → przeglądarka z nawigacją; reszta → plik wprost
                    ext = item.suffix.lower()
                    href = (f'{q}?view=1'
                            if (ext in IMG_EXT or ext in AUDIO_EXT
                                or ext in ('.md', '.docx', '.doc'))
                            else q)
                    lek = ('<a href="#" class="dl lek" data-n="' + q
                           + '" title="Lektor → audio">🔊</a>'
                           if ext in ('.md', '.docx', '.txt') else '')
                    rows.append(
                        f'<tr><td><a href="{q}?dl=1" class="dl" title="Pobierz">⬇</a>'
                        f'<a href="#" class="dl del" data-n="{q}" title="Usuń (do .kosz)">🗑</a>'
                        f'<a href="#" class="dl ren" data-n="{q}" title="Zmień nazwę">✎</a>'
                        f'<a href="#" class="dl cpy" data-n="{q}" title="Kopiuj nazwę">⧉</a>'
                        f'{lek}'
                        f'<a href="{href}">{icon} {item.name}</a></td>'
                        f'<td data-sort="{s}">{_fmt_size(s)}</td>'
                        f'<td data-sort="{st.st_mtime:.0f}">{mt_s}</td>'
                        f'<td data-sort="{bt:.0f}">{bt_s}</td></tr>')
                n += 1
            except Exception:
                pass

        html = [
            '<!doctype html><meta charset=utf-8>',
            '<meta name="viewport" content="width=device-width,initial-scale=1">',
            f'<title>{rel_url}</title>',
            f'<style>{STYLE}</style>',
            f'<h2>{breadcrumb}</h2>',
            f'<p class="muted">{n} pozycji · kliknij nagłówek aby sortować · '
            f'⟳ auto-odświeżanie · <a href="/?lektorq=1">🔊 kolejka lektora</a>'
            f'{_battery_html()}</p>',
            '<table><thead><tr>'
            '<th>Nazwa</th><th>Rozmiar</th><th>Modyfikacja</th><th>Utworzono</th>'
            '</tr></thead><tbody>',
            *rows,
            '</tbody></table>',
            SORT_JS,
            DROP_JS,
            DEL_JS,
        ]
        return web.Response(text='\n'.join(html), content_type='text/html')

    # plik — ?dl=1 wymusza pobieranie, ?view=1 dla zdjęć otwiera przeglądarkę
    return await _serve_file(request, target)


async def delete_item(request):
    """DELETE na pliku = przeniesienie do ROOT/.kosz/ (nic nie znika trwale).
    Katalogi: tylko puste (rmdir). .kosz ukryty w listingu (dotfile)."""
    raw = request.match_info.get('path', '').strip('/')
    try:
        target = (ROOT / raw).resolve()
    except Exception:
        return web.Response(status=400)
    if ROOT not in target.parents:          # ROOT samego nie ruszamy
        return web.Response(status=403)
    if not target.exists():
        return web.Response(status=404)
    if target.is_dir():
        if any(target.iterdir()):
            return web.Response(status=400, text='Katalog niepusty')
        target.rmdir()
        return web.json_response({'deleted': raw})
    trash = ROOT / '.kosz'
    trash.mkdir(exist_ok=True)
    ts = datetime.now().strftime('%Y%m%d_%H%M%S')
    dest = trash / f'{ts}_{target.name}'
    target.rename(dest)
    return web.json_response({'deleted': raw, 'kosz': dest.name})


CZYTAJ_TTS = '/mnt/data/sprawozdania/EXPORT/czytaj_tts.py'
LEKTOR_CONF = '/mnt/data/sprawozdania/EXPORT/lektor-ustawienia.conf'
_LEKTOR_LOCK = None        # asyncio.Lock tworzony leniwie (jeden lektor naraz)
_LEKTOR_SEQ = 0
_LEKTOR_QUEUE: list = []   # rejestr zadań: id/out/src/fmt/state/cancelled/proc
_LEKTOR_PAUSED = False     # pauza całej kolejki (⏸ w widoku kolejki)
_LEKTOR_SHUTDOWN = False   # „wyłącz konsolę po ukończeniu kolejki" (⏻)


LEKTOR_QFILE = Path('/mnt/data/lektor_queue.json')


def _lektor_save_queue():
    """Trwałość kolejki: przeżywa restart serwera i REBOOT konsoli
    (na starcie zadania wracają; czytaj_tts wznawia z checkpointu chunków)."""
    import json
    try:
        LEKTOR_QFILE.write_text(json.dumps({
            'paused': _LEKTOR_PAUSED,
            'shutdown': _LEKTOR_SHUTDOWN,
            'jobs': [{'src': j['src'], 'out': j['out'], 'fmt': j['fmt'],
                      'opisy': j.get('opisy', '')}
                     for j in _LEKTOR_QUEUE if not j['cancelled']]},
            ensure_ascii=False))
    except Exception:
        pass


async def _lektor_restore(app):
    """on_startup: odtwórz kolejkę z dysku. Zadanie, które właśnie generuje
    osierocony proces (KillMode=process), pomijamy — dokończy się samo."""
    import json
    global _LEKTOR_PAUSED, _LEKTOR_SHUTDOWN
    try:
        data = json.loads(LEKTOR_QFILE.read_text())
    except Exception:
        return
    _LEKTOR_PAUSED = bool(data.get('paused'))
    _LEKTOR_SHUTDOWN = bool(data.get('shutdown'))
    asyncio.ensure_future(_lektor_shutdown_watch())
    prog = _lektor_progress()
    busy_stem = Path(prog['out']).stem if prog else None
    for it in data.get('jobs', []):
        if not Path(it['src']).exists():
            continue
        if busy_stem and Path(it['out']).stem == busy_stem \
                and _ext_lektor_running():
            continue   # już generowane przez osierocony/zewnętrzny proces
        _lektor_new_job(it['src'], Path(it['out']), it['fmt'],
                        it.get('opisy', ''))


def _lektor_new_job(src, out, fmt, opisy='') -> dict:
    """Rejestracja zadania + start workera (wspólne dla 🔊 i „przejdź
    do następnego" przy pauzie)."""
    global _LEKTOR_SEQ, _LEKTOR_LOCK
    if _LEKTOR_LOCK is None:
        _LEKTOR_LOCK = asyncio.Lock()
    _LEKTOR_SEQ += 1
    import time as _t
    job = {'id': _LEKTOR_SEQ, 'out': str(out), 'src': str(src), 'fmt': fmt,
           'opisy': opisy, 'state': 'queued', 'cancelled': False,
           'proc': None, 'started': _t.time()}
    _LEKTOR_QUEUE.append(job)
    _lektor_save_queue()
    asyncio.ensure_future(_lektor_run(job))
    return job


async def _lektor_run(job: dict):
    out = Path(job['out'])
    shutdown = False
    try:
        async with _LEKTOR_LOCK:
            if job['cancelled']:
                return
            # czekaj: pauza kolejki ORAZ lektor spoza serwera (agent/CLI)
            while _LEKTOR_PAUSED or _ext_lektor_running():
                await asyncio.sleep(5)
                if job['cancelled']:
                    return
            job['state'] = 'running'
            import time as _t
            job['started'] = _t.time()
            # stderr do logu — bez tego pad lektora był niemy (incydent:
            # zadanie znikało z kolejki bez śladu i bez pliku)
            errlog = open('/mnt/data/lektor_errors.log', 'ab')
            errlog.write(f'\n=== {Path(job["out"]).name} ===\n'.encode())
            cmd = ['python3', CZYTAJ_TTS, job['src'], '-o', job['out'],
                   '--format', job['fmt']]
            if job.get('opisy') in ('tak', 'nie'):
                cmd += ['--opisy', job['opisy']]
            proc = await asyncio.create_subprocess_exec(
                *cmd,
                stdout=asyncio.subprocess.DEVNULL,
                stderr=errlog)
            job['proc'] = proc
            await proc.wait()
            errlog.close()
            if proc.returncode not in (0, None) and not job['cancelled']:
                job['state'] = 'failed'
            if job['cancelled']:
                # przerwane — sprzątnij TYLKO pliki powstałe w trakcie
                # TEGO zadania (mtime > start)
                for suf in ('.mp3', '.wav', '.flac'):
                    p = out.with_suffix(suf)
                    try:
                        if p.stat().st_mtime >= job['started']:
                            p.unlink()
                    except FileNotFoundError:
                        pass
    except asyncio.CancelledError:
        # shutdown/restart serwera — zadanie ma PRZETRWAĆ w pliku kolejki
        # (restore odtworzy je na starcie); nie nadpisujemy rejestru
        shutdown = True
        raise
    finally:
        try:
            _LEKTOR_QUEUE.remove(job)
        except ValueError:
            pass
        if not shutdown:
            _lektor_save_queue()
            _lektor_maybe_shutdown()


def _ext_lektor_pids() -> set:
    """PID-y czytaj_tts.py uruchomione POZA serwerem (agent z Discorda,
    CLI) — pgrep minus nasze własne dzieci z kolejki."""
    import subprocess
    r = subprocess.run(['pgrep', '-f', 'czytaj_tts.py'],
                       capture_output=True, text=True)
    pids = {int(x) for x in r.stdout.split()} if r.returncode == 0 else set()
    # tylko realne interpretery Pythona — pgrep -f łapie też powłoki/wrappery,
    # których CMDLINE zawiera nazwę skryptu (np. sesję ssh agenta!)
    real = set()
    for p in pids:
        try:
            comm = Path(f'/proc/{p}/comm').read_text().strip()
        except OSError:
            continue
        if comm.startswith('python'):
            real.add(p)
    ours = {j['proc'].pid for j in _LEKTOR_QUEUE
            if j.get('proc') is not None and j['proc'].returncode is None}
    return real - ours


def _ext_lektor_running() -> bool:
    return bool(_ext_lektor_pids())


def _lektor_busy() -> bool:
    global _LEKTOR_LOCK
    return bool(_LEKTOR_QUEUE) or _ext_lektor_running() \
        or (_LEKTOR_LOCK is not None and _LEKTOR_LOCK.locked())


def _lektor_progress():
    """Postęp bieżącej generacji (pisze czytaj_tts.py; lektor jest jeden,
    więc plik globalny wystarcza). None = brak danych."""
    import json
    import time as _t
    try:
        p = Path('/tmp/lektor_progress.json')
        if _t.time() - p.stat().st_mtime > 300:   # stęchły = po crashu
            return None
        return json.loads(p.read_text())
    except Exception:
        return None


def _lektor_queue_json() -> dict:
    prog = _lektor_progress()
    jobs = []
    for j in _LEKTOR_QUEUE:
        if j['cancelled']:
            continue
        e = {'id': j['id'], 'out': Path(j['out']).name,
             'src': str(Path(j['src']).relative_to(ROOT))
             if str(j['src']).startswith(str(ROOT)) else Path(j['src']).name,
             'fmt': j['fmt'], 'state': j['state']}
        if j['state'] == 'running' and prog:
            e['pct'] = prog.get('pct', 0)
            e['chunk'] = f"{prog.get('chunk', 0)}/{prog.get('chunks', 0)}"
        jobs.append(e)
    ext = _ext_lektor_running()
    out = {'jobs': jobs, 'external': ext, 'paused': _LEKTOR_PAUSED,
           'shutdown': _LEKTOR_SHUTDOWN, 'bat': _battery_html()}
    if ext and prog and not any(j['state'] == 'running' for j in jobs):
        out['ext_pct'] = prog.get('pct', 0)
        out['ext_chunk'] = f"{prog.get('chunk', 0)}/{prog.get('chunks', 0)}"
        out['ext_out'] = prog.get('out', '')
    return out


LEKTORQ_PAGE = (
    '<!doctype html><meta charset=utf-8>'
    '<meta name="viewport" content="width=device-width,initial-scale=1">'
    '<title>Kolejka lektora</title><style>'
    # styl 1:1 z listingiem serwera plików (jasny motyw)
    'body{font-family:system-ui,sans-serif;max-width:1000px;margin:1.2em auto;'
    'padding:0 1em;background:#f7f7f9;color:#222}'
    'h2{color:#333;font-size:1.1em;word-break:break-all}'
    'h2 a{text-decoration:none}h2 a:hover{text-decoration:underline}'
    '.sep{color:#999}'
    '.muted{color:#888;font-size:.85em}'
    'table{width:100%;border-collapse:collapse;background:#fff;'
    'border-radius:6px;overflow:hidden;box-shadow:0 1px 3px rgba(0,0,0,.08)}'
    'th,td{text-align:left;padding:.5em .7em;border-bottom:1px solid #eee;'
    'font-size:.92em}'
    'th{background:#eef2f7;user-select:none}'
    'tr:hover td{background:#f4f7fb}'
    '.run{color:#1d7a36;font-weight:600}.que{color:#9a7b00}'
    '.ext{color:#888;font-style:italic}'
    '.x{color:#c00;cursor:pointer;text-decoration:none;font-weight:700;'
    'opacity:.55}.x:hover{opacity:1}'
    '.empty{text-align:center;color:#888;padding:2em}'
    '.pb{width:130px;height:13px;background:#e4e8ee;border-radius:7px;'
    'overflow:hidden;display:inline-block;vertical-align:middle;'
    'margin-right:.5em}'
    '.pb>i{display:block;height:100%;background:linear-gradient(90deg,'
    '#3584e4,#33a02c);transition:width .8s}'
    '.pct{font-size:.85em;color:#888}'
    # mobile: bez kolumny źródła, duże pola dotyku dla ⏸/✕
    '@media (max-width:700px){'
    'body{margin:.6em auto;padding:0 .5em}'
    'th:nth-child(3),td:nth-child(3){display:none}'
    'th,td{padding:.8em .5em}'
    '.x,.p{font-size:1.3em;padding:.2em .3em}'
    '.pb{width:90px}'
    '}'
    '</style>'
    '<h2><a href="/" class="dir">📁 root</a> <span class="sep">/</span> '
    '🔊 Kolejka lektora</h2>'
    '<p class="muted">podgląd na żywo (co 3 s) · ✕ usuwa pozycję '
    '(trwająca generacja zostaje przerwana) · '
    '<a id="shd" href="#" style="text-decoration:none"></a>'
    ' · <span id="st"></span></p>'
    '<table><thead><tr><th>#</th><th>plik wynikowy</th><th>źródło</th>'
    '<th>format</th><th>stan</th><th></th></tr></thead>'
    '<tbody id="tb"></tbody></table>'
    '<script>'
    'async function load(){'
    'try{const j=await(await fetch("/?lektorqj=1",{cache:"no-store"})).json();'
    'const tb=document.getElementById("tb");let h="";let i=0;'
    'function bar(p,c){return p==null?"GENERUJE":'
    '"<span class=pb><i style=\'width:"+p+"%\'></i></span>'
    '<span class=pct>"+p+"% ("+(c||"")+")</span>";}'
    'function acts(id,run){return "<td style=\'white-space:nowrap\'>"'
    '+(run?"<a class=p data-i="+id+" href=# title=\'Pauza\' '
    'style=\'margin-right:.6em;text-decoration:none\'>⏸</a>":"")'
    '+"<a class=x data-i="+id+" href=#>✕</a></td>";}'
    'if(j.paused){h+="<tr><td colspan=6 style=\'background:#fff7e0;'
    'color:#7a5b00;font-weight:600;text-align:center;padding:.7em\'>'
    '⏸ KOLEJKA WSTRZYMANA &nbsp; '
    '<a id=res href=# style=\'color:#1a5fb4\'>▶ wznów</a></td></tr>";}'
    'if(j.external){h+="<tr><td>—</td><td colspan=3>"'
    '+(j.ext_out?j.ext_out+" <span class=ext>(proces niezależny — '
    'zlecenie agenta albo kontynuacja po restarcie serwera)</span>"'
    ':"<span class=ext>lektor uruchomiony poza serwerem '
    '(agent Discord / CLI)</span>")'
    '+"</td><td class="+(j.paused?"que":"run")+">"'
    '+(j.paused?"⏸ wstrzymane":bar(j.ext_pct,j.ext_chunk))+"</td>"'
    '+acts("ext",!j.paused)+"</tr>";}'
    'for(const x of j.jobs){i++;'
    'const run=(x.state==="running");'
    'h+="<tr><td>"+i+"</td><td>"+x.out+"</td><td style=\'color:#888\'>"'
    '+x.src+"</td><td>"+x.fmt+"</td><td class="'
    '+(run?"run":"que")+">"'
    '+(x.state==="paused"?"⏸ wstrzymane":(run?bar(x.pct,x.chunk):"czeka"))'
    '+"</td>"+acts(x.id,run)+"</tr>";}'
    'if(!h)h="<tr><td colspan=6 class=empty>Kolejka pusta — lektor wolny</td></tr>";'
    'tb.innerHTML=h;'
    'document.getElementById("st").innerHTML='
    '"odświeżono "+new Date().toLocaleTimeString()+(j.bat||"");'
    'const sh=document.getElementById("shd");'
    'sh.dataset.on=j.shutdown?"1":"0";'
    'sh.innerHTML=j.shutdown'
    '?"⏻ wyłącz konsolę po ukończeniu: <b style=\'color:#1d7a36\'>WŁĄCZONE</b>"'
    ':"⏻ wyłącz konsolę po ukończeniu: <b style=\'color:#888\'>wyłączone</b>";'
    '}catch(e){}}'
    # pauza: pytanie co dalej — następny plik czy wstrzymanie całej kolejki
    'function askPause(ext){return new Promise(res=>{'
    'const ov=document.createElement("div");'
    'ov.style.cssText="position:fixed;inset:0;background:rgba(0,0,0,.5);'
    'display:flex;align-items:center;justify-content:center;z-index:99";'
    'ov.innerHTML=\'<div style="background:#fff;border-radius:10px;'
    'padding:1.1em 1.4em;max-width:92vw;box-shadow:0 8px 30px rgba(0,0,0,.35)">'
    '<div style="font-weight:600;margin-bottom:.5em">⏸ Pauza generowania</div>'
    '<div style="color:#556;font-size:.9em;margin-bottom:1em">'
    'Co zrobić z bieżącym plikiem?</div>'
    '<div style="display:flex;flex-direction:column;gap:.5em">\'+'
    '(ext?"":\'<button data-m="skip">⏭ Przejdź do następnego pliku '
    '(ten wróci na koniec kolejki)</button>\')+'
    '\'<button data-m="hold">⏸ Wstrzymaj CAŁĄ kolejkę '
    '(wznowisz przyciskiem ▶)</button>'
    '<button data-m="x">Anuluj</button></div></div>\';'
    'ov.querySelectorAll("button").forEach(b=>{'
    'b.style.cssText="padding:.55em 1em;border:1px solid #b8c0cc;'
    'border-radius:6px;background:#f2f5f9;cursor:pointer;font-size:.95em;'
    'text-align:left";'
    'b.onclick=()=>{ov.remove();res(b.dataset.m==="x"?null:b.dataset.m);};});'
    'ov.onclick=e=>{if(e.target===ov){ov.remove();res(null);}};'
    'document.body.appendChild(ov);});}'
    'document.addEventListener("click",async e=>{'
    'const s=e.target.closest("a#shd");'
    'if(s){e.preventDefault();'
    'const to=s.dataset.on==="1"?"0":"1";'
    'if(to==="1"&&!confirm("Konsola WYŁĄCZY SIĘ automatycznie po '
    'ukończeniu wszystkich pozycji kolejki (1 min na anulowanie). '
    'Włączyć?"))return;'
    'let r=await fetch("/?lektorqshutdown="+to,{method:"POST"});'
    'let j=await r.json().catch(()=>({}));'
    'if(j.status==="empty-queue"){'
    'if(confirm("UWAGA: kolejka jest PUSTA — konsola wyłączy się '
    'JUŻ ZA MINUTĘ, nie po przyszłych zadaniach.\\n\\n'
    'Na pewno wyłączyć konsolę teraz?"))'
    'await fetch("/?lektorqshutdown=1&force=1",{method:"POST"});}'
    'load();return;}'
    'const r=e.target.closest("a#res");'
    'if(r){e.preventDefault();'
    'await fetch("/?lektorqresume=1",{method:"POST"});load();return;}'
    'const p=e.target.closest("a.p");'
    'if(p){e.preventDefault();'
    'const m=await askPause(p.dataset.i==="ext");'
    'if(!m)return;'
    'await fetch("/?lektorqpause="+p.dataset.i+"&mode="+m,{method:"POST"});'
    'load();return;}'
    'const a=e.target.closest("a.x");if(!a)return;e.preventDefault();'
    'if(!confirm("Usunąć tę pozycję z kolejki lektora?"+'
    '"\\n(trwająca generacja zostanie przerwana)"))return;'
    'await fetch("/?lektorqdel="+a.dataset.i,{method:"POST"});load();});'
    'load();setInterval(load,3000);'
    '</script>')


def _lektor_fmt() -> str:
    """Format z lektor-ustawienia.conf (mp3|wav|flac), domyślnie mp3."""
    try:
        for line in Path(LEKTOR_CONF).read_text().splitlines():
            line = line.split('#', 1)[0]
            if '=' in line:
                k, v = line.split('=', 1)
                if k.strip() == 'format' and v.strip() in ('mp3', 'wav', 'flac'):
                    return v.strip()
    except Exception:
        pass
    return 'mp3'


def _docx_to_txt(p: Path) -> Path:
    """Awaryjne źródło dla lektora: tekst wprost z DOCX (python-docx)."""
    import docx
    d = docx.Document(str(p))
    parts = [par.text for par in d.paragraphs]
    for tbl in d.tables:
        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(' . '.join(cells))
    tmp = Path('/tmp') / (p.stem + '_lektor_src.txt')
    tmp.write_text('\n'.join(parts), encoding='utf-8')
    return tmp


async def print_item(request):
    """POST ?print=1 na .md/.docx/.pdf — druk na Canon G3070 (CUPS).
    .md → export_to_docx → soffice→PDF (paginacja jak w Wordzie) → lp.
    Async; błędy → /mnt/data/print_errors.log."""
    raw = request.match_info.get('path', '').strip('/')
    try:
        target = (ROOT / raw).resolve()
    except Exception:
        return web.Response(status=400)
    if ROOT not in target.parents or not target.is_file():
        return web.Response(status=403)
    ext = target.suffix.lower()
    if ext not in ('.md', '.docx', '.pdf'):
        return web.Response(status=400, text='Druk: .md/.docx/.pdf')

    async def _run():
        import time as _t
        log = open('/mnt/data/print_errors.log', 'ab')
        log.write(f'\n=== {_t.strftime("%F %T")} {target.name} ===\n'.encode())

        async def sh(*cmd):
            p = await asyncio.create_subprocess_exec(
                *cmd, stdout=log, stderr=log)
            await p.wait()
            return p.returncode

        try:
            pdf = target
            if ext == '.md':
                tmp_docx = Path('/tmp') / (target.stem + '_print.docx')
                if await sh('python3',
                            '/mnt/data/sprawozdania/EXPORT/export_to_docx.py',
                            str(target), '-o', str(tmp_docx)):
                    return
                pdf = Path('/tmp') / (tmp_docx.stem + '.pdf')
                if await sh('soffice', '--headless',
                            '-env:UserInstallation=file:///tmp/lo_print',
                            '--convert-to', 'pdf', '--outdir', '/tmp',
                            str(tmp_docx)):
                    return
            elif ext == '.docx':
                pdf = Path('/tmp') / (target.stem + '.pdf')
                if await sh('soffice', '--headless',
                            '-env:UserInstallation=file:///tmp/lo_print',
                            '--convert-to', 'pdf', '--outdir', '/tmp',
                            str(target)):
                    return
            await sh('lp', '-d', 'Canon_G3070', str(pdf))
        finally:
            log.close()

    asyncio.ensure_future(_run())
    return web.json_response({'status': 'wysłano do druku'}, status=202)


async def lektor_item(request):
    """POST ?lektor=1 na .md/.docx/.txt — generacja audio w TLE
    (czytaj_tts.py). Wynik: <nazwa>_lektor.<fmt> obok pliku (fmt z conf);
    auto-odświeżanie listingu pokaże go po zakończeniu."""
    raw = request.match_info.get('path', '').strip('/')
    try:
        target = (ROOT / raw).resolve()
    except Exception:
        return web.Response(status=400)
    if ROOT not in target.parents or not target.is_file():
        return web.Response(status=403)
    ext = target.suffix.lower()
    if ext not in ('.md', '.docx', '.txt'):
        return web.Response(status=400, text='Lektor czyta .md/.docx/.txt')

    # źródło tekstu: .md wprost; .docx → bliźniaczy .md (ten sam stem,
    # ten katalog lub ../processed/) albo ekstrakcja tekstu z DOCX
    src = target
    if ext == '.docx':
        for cand in (target.parent / (target.stem + '.md'),
                     target.parent.parent / 'processed' / (target.stem + '.md')):
            if cand.exists():
                src = cand
                break
        else:
            try:
                src = _docx_to_txt(target)
            except Exception as e:
                return web.Response(status=500, text=f'Ekstrakcja DOCX: {e}')

    fmt = request.query.get('fmt', '').lower() or _lektor_fmt()
    if fmt not in ('mp3', 'wav', 'flac'):
        fmt = 'mp3'
    # konwencja sprawozdań: źródło w processed/ → audio do exports/ obok
    # (ta sama lokalizacja co lektor agenta z Discorda); inaczej obok pliku
    out_dir = target.parent
    if out_dir.name == 'processed' and (out_dir.parent / 'exports').is_dir():
        out_dir = out_dir.parent / 'exports'
    out = out_dir / f'{target.stem}_lektor.{fmt}'
    if any(j['out'] == str(out) and not j['cancelled'] for j in _LEKTOR_QUEUE):
        return web.json_response({'status': 'duplikat', 'out': out.name})

    # lektor zajęty (przeglądarka ALBO agent z Discorda) → bez flagi queue
    # zwracamy 'busy'; UI pyta usera o dopisanie do kolejki
    if _lektor_busy() and 'queue' not in request.query:
        return web.json_response(
            {'status': 'busy', 'pending': len(_LEKTOR_QUEUE)})

    queued = bool(_LEKTOR_QUEUE) or _ext_lektor_running() or _LEKTOR_PAUSED
    opisy = request.query.get('opisy', '')
    if opisy not in ('tak', 'nie'):
        opisy = ''
    _lektor_new_job(src, out, fmt, opisy)
    return web.json_response(
        {'status': 'queued' if queued else 'start',
         'out': out.name, 'position': len(_LEKTOR_QUEUE)}, status=202)


async def rename_item(request):
    """POST ?rename=<nowa-nazwa> na pliku/katalogu = zmiana nazwy (ten sam
    katalog, bez nadpisywania — 409 gdy cel istnieje)."""
    raw = request.match_info.get('path', '').strip('/')
    try:
        target = (ROOT / raw).resolve()
    except Exception:
        return web.Response(status=400)
    if ROOT not in target.parents:
        return web.Response(status=403)
    if not target.exists():
        return web.Response(status=404)
    new_name = Path(request.query.get('rename', '')).name.strip()
    if not new_name or new_name.startswith('.'):
        return web.Response(status=400, text='Nieprawidłowa nazwa')
    dest = target.parent / new_name
    if dest.exists():
        return web.Response(status=409, text='Plik o tej nazwie już istnieje')
    target.rename(dest)
    return web.json_response({'renamed': target.name, 'to': new_name})


async def lektor_pause(request):
    """POST ?lektorqpause=<id|ext>&mode=hold|skip
    hold = SIGSTOP bieżącej generacji + wstrzymanie CAŁEJ kolejki;
    skip = przerwij bieżącą, jej zadanie wraca NA KONIEC kolejki,
           startuje następny plik."""
    global _LEKTOR_PAUSED
    import os
    import signal as _sig
    rid = request.query.get('lektorqpause', '')
    mode = request.query.get('mode', 'hold')
    if rid == 'ext':
        # zewnętrzny lektor (agent/CLI): tylko hold/zamrożenie
        for pid in _ext_lektor_pids():
            try:
                os.kill(pid, _sig.SIGSTOP)
            except ProcessLookupError:
                pass
        _LEKTOR_PAUSED = True
        _lektor_save_queue()
        return web.json_response({'paused': 'ext'})
    try:
        jid = int(rid)
    except ValueError:
        return web.Response(status=400)
    job = next((j for j in _LEKTOR_QUEUE if j['id'] == jid), None)
    if job is None:
        return web.Response(status=404)
    if mode == 'skip':
        # bieżący na koniec kolejki, następny rusza
        job['cancelled'] = True
        if job.get('proc') is not None and job['state'] == 'running':
            try:
                job['proc'].kill()
            except ProcessLookupError:
                pass
        nj = _lektor_new_job(job['src'], Path(job['out']), job['fmt'],
                             job.get('opisy', ''))
        return web.json_response({'skipped': jid, 'requeued_as': nj['id']})
    # hold
    if job.get('proc') is not None and job['state'] == 'running':
        try:
            os.kill(job['proc'].pid, _sig.SIGSTOP)
            job['state'] = 'paused'
        except ProcessLookupError:
            pass
    _LEKTOR_PAUSED = True
    _lektor_save_queue()
    return web.json_response({'paused': jid})


async def lektor_resume(request):
    """POST ?lektorqresume=1 — wznowienie kolejki (SIGCONT zamrożonych)."""
    global _LEKTOR_PAUSED
    import os
    import signal as _sig
    _LEKTOR_PAUSED = False
    _lektor_save_queue()
    for j in _LEKTOR_QUEUE:
        if j['state'] == 'paused' and j.get('proc') is not None:
            try:
                os.kill(j['proc'].pid, _sig.SIGCONT)
                j['state'] = 'running'
            except ProcessLookupError:
                pass
    for pid in _ext_lektor_pids():
        try:
            os.kill(pid, _sig.SIGCONT)
        except ProcessLookupError:
            pass
    return web.json_response({'resumed': True})


async def lektor_cancel(request):
    """POST ?lektorqdel=<id|ext> — usunięcie pozycji z kolejki lektora;
    trwająca generacja zostaje ubita (i sprzątnięte częściowe pliki).
    'ext' = przerwij lektora uruchomionego poza serwerem (agent/CLI)."""
    raw_id = request.query.get('lektorqdel', '')
    if raw_id == 'ext':
        import signal
        killed = []
        for pid in _ext_lektor_pids():
            try:
                import os
                os.kill(pid, signal.SIGKILL)
                killed.append(pid)
            except ProcessLookupError:
                pass
        return web.json_response({'cancelled': 'ext', 'pids': killed})
    try:
        jid = int(raw_id)
    except ValueError:
        return web.Response(status=400)
    for j in _LEKTOR_QUEUE:
        if j['id'] == jid:
            j['cancelled'] = True
            if j['state'] == 'running' and j.get('proc') is not None:
                try:
                    j['proc'].kill()
                except ProcessLookupError:
                    pass
            else:
                try:
                    _LEKTOR_QUEUE.remove(j)
                except ValueError:
                    pass
            return web.json_response({'cancelled': jid})
    return web.Response(status=404)


async def lektor_shutdown_toggle(request):
    """POST ?lektorqshutdown=1|0 — wyłączenie konsoli po ukończeniu kolejki.
    0 anuluje też zaplanowane już `shutdown` (gdy kolejka właśnie się
    skończyła i odliczanie trwa)."""
    global _LEKTOR_SHUTDOWN
    import subprocess
    val = request.query.get('lektorqshutdown', '0') == '1'
    # PUSTA kolejka + włączenie = natychmiastowe wyłączenie konsoli —
    # wymagaj jawnego force (incydent 2026-06-07 11:15: konsola zgasła
    # w trakcie odtwarzania muzyki)
    if val and not _LEKTOR_QUEUE and not _ext_lektor_running() \
            and 'force' not in request.query:
        return web.json_response({'status': 'empty-queue'})
    _LEKTOR_SHUTDOWN = val
    _lektor_save_queue()
    if not val:
        subprocess.run(['shutdown', '-c'], capture_output=True)
    else:
        _lektor_maybe_shutdown()
    return web.json_response({'shutdown': _LEKTOR_SHUTDOWN})


def _lektor_maybe_shutdown():
    """Kolejka pusta + nic nie generuje + flaga ⏻ → shutdown za 1 min
    (okno na anulowanie togglem); flaga konsumowana."""
    global _LEKTOR_SHUTDOWN
    if not _LEKTOR_SHUTDOWN:
        return
    if _LEKTOR_QUEUE or _ext_lektor_running():
        return
    import subprocess
    subprocess.run(['shutdown', '-h', '+1',
                    'Lektor ukończył kolejkę — wyłączanie'],
                   capture_output=True)
    _LEKTOR_SHUTDOWN = False
    _lektor_save_queue()


async def _lektor_shutdown_watch():
    """Strażnik flagi ⏻ dla zadań ZEWNĘTRZNYCH (agent) — _lektor_run ich
    nie widzi, więc sprawdzamy co 60 s."""
    while True:
        await asyncio.sleep(60)
        if _LEKTOR_SHUTDOWN:
            _lektor_maybe_shutdown()


async def upload(request):
    if 'print' in request.query:
        return await print_item(request)
    if 'lektorqshutdown' in request.query:
        return await lektor_shutdown_toggle(request)
    if 'lektorqpause' in request.query:
        return await lektor_pause(request)
    if 'lektorqresume' in request.query:
        return await lektor_resume(request)
    if 'lektorqdel' in request.query:
        return await lektor_cancel(request)
    if 'rename' in request.query:
        return await rename_item(request)
    if 'lektor' in request.query:
        return await lektor_item(request)
    """POST multipart na katalog = wgranie plików (drag&drop z przeglądarki).
    Duplikaty nazw dostają sufiks z timestampem (jak w bocie) — nic nie nadpisujemy."""
    raw = request.match_info.get('path', '').strip('/')
    try:
        target = (ROOT / raw).resolve()
    except Exception:
        return web.Response(status=400)
    if ROOT not in target.parents and target != ROOT:
        return web.Response(status=403)
    if not target.is_dir():
        return web.Response(status=400, text='Cel nie jest katalogiem')
    saved = []
    reader = await request.multipart()
    async for part in reader:
        if part.name != 'file' or not part.filename:
            continue
        name = Path(part.filename).name
        if not name or name.startswith('.'):
            continue
        dest = target / name
        if dest.exists():
            ts = datetime.now().strftime('%Y%m%d_%H%M%S')
            dest = target / f'{dest.stem}_{ts}{dest.suffix}'
        with open(dest, 'wb') as f:
            while True:
                chunk = await part.read_chunk(64 * 1024)
                if not chunk:
                    break
                f.write(chunk)
        saved.append(dest.name)
    return web.json_response({'saved': saved})


async def _serve_file(request, target):
    if 'dl' in request.query:
        return web.FileResponse(target, headers={
            'Cache-Control': 'no-cache',
            'Content-Disposition': f"attachment; filename*=UTF-8''{quote(target.name)}"})

    if 'view' in request.query and target.suffix.lower() in AUDIO_EXT:
        return web.Response(text=render_audio_page(target),
                            content_type='text/html')

    # audio bez ?view: poprawny MIME (mimetypes nie zna .flac → pobieranie
    # zamiast odtwarzania w <audio>)
    if target.suffix.lower() in AUDIO_EXT and 'dl' not in request.query:
        return web.FileResponse(target, headers={
            'Content-Type': AUDIO_MIME[target.suffix.lower()]})

    # surowy PDF z konwersji (źródło dla <embed> w stronie podglądu)
    if 'pdf' in request.query and target.suffix.lower() in ('.docx', '.doc'):
        pdf = await docx_to_pdf(target)
        if pdf is None:
            return web.Response(status=500, text='Konwersja DOCX nie powiodla sie')
        return web.FileResponse(pdf, headers={
            'Content-Type': 'application/pdf',
            'Content-Disposition': f"inline; filename*=UTF-8''{quote(target.stem)}.pdf",
            'Cache-Control': 'no-cache'})

    if 'view' in request.query and target.suffix.lower() in ('.docx', '.doc'):
        # strona-opakowanie: PDF w <embed> + poll mtime DOCX co 3 s —
        # po zmianie pliku (np. re-eksport przez agenta) podgląd sam się
        # przeładowuje (świeża konwersja, cache-bust parametrem v=)
        q = quote(target.name)
        mt = target.stat().st_mtime
        page = (
            '<!doctype html><meta charset=utf-8>'
            '<meta name="viewport" content="width=device-width,initial-scale=1">'
            f'<title>{target.name}</title>'
            '<style>body{margin:0;height:100vh;display:flex;'
            'flex-direction:column}'
            '.bar{display:flex;gap:.6em;align-items:center;padding:.4em .9em;'
            'background:#26292f;color:#ddd;font-family:system-ui,sans-serif;'
            'font-size:.9em}'
            '.bar a{color:#7ab7ff;text-decoration:none;border:1px solid '
            '#3a3f47;border-radius:5px;padding:.2em .6em}'
            'embed{flex:1;width:100%;border:0}</style>'
            f'<div class="bar"><a href="./">📁 folder</a>'
            f'<span style="word-break:break-all">{target.name}</span>'
            f'<span id="st" style="color:#8a93a0"></span>'
            f'<a href="{q}?dl=1" style="margin-left:auto">⬇ DOCX</a></div>'
            f'<embed id="pv" src="{q}?pdf=1&v={mt}" '
            'type="application/pdf">'
            '<script>(function(){'
            f'let mt={mt};'
            'async function chk(){try{'
            f'const j=await(await fetch("{q}?mt=1",'
            '{cache:"no-store"})).json();'
            'if(j.mt!==mt){mt=j.mt;'
            'document.getElementById("st").textContent="odświeżam...";'
            f'document.getElementById("pv").src="{q}?pdf=1&v="+mt;'
            'setTimeout(()=>document.getElementById("st").textContent="",4000);}'
            '}catch(e){}}'
            'setInterval(chk,3000);})();</script>')
        return web.Response(text=page, content_type='text/html')

    if 'mt' in request.query:
        return web.json_response({'mt': target.stat().st_mtime})

    if 'view' in request.query and target.suffix.lower() == '.md':
        return web.Response(text=render_md_page(target), content_type='text/html')

    if 'view' in request.query and target.suffix.lower() in IMG_EXT:
        siblings = sorted(
            (p for p in target.parent.iterdir()
             if p.is_file() and p.suffix.lower() in IMG_EXT
             and not p.name.startswith('.')),
            key=lambda p: p.name.lower())
        names = [p.name for p in siblings]
        idx = names.index(target.name) if target.name in names else 0
        prv = quote(names[idx-1]) + '?view=1' if idx > 0 else None
        nxt = quote(names[idx+1]) + '?view=1' if idx < len(names)-1 else None
        a_prev = (f'<a href="{prv}" id="prev">← poprzednie</a>' if prv
                  else '<a class="nav-off">← poprzednie</a>')
        a_next = (f'<a href="{nxt}" id="next">następne →</a>' if nxt
                  else '<a class="nav-off">następne →</a>')
        html = (
            '<!doctype html><meta charset=utf-8>'
            '<meta name="viewport" content="width=device-width,initial-scale=1">'
            f'<title>{target.name}</title><style>{VIEWER_STYLE}</style>'
            f'<div class="bar"><a href="./">📁 folder</a>{a_prev}{a_next}'
            f'<span class="name">{target.name}</span>'
            f'<span class="cnt">{idx+1} / {len(names)}</span>'
            f'<span class="cnt" id="zl">100%</span>'
            f'<a href="#" id="fitb" title="Dopasuj do okna (dwuklik / klawisz 0)">⊡ dopasuj</a>'
            f'<a href="#" id="natb" title="Rozmiar naturalny (klawisz 1)">1:1</a>'
            f'<a href="{quote(target.name)}?dl=1">⬇ pobierz</a></div>'
            f'<div class="stage"><img src="{quote(target.name)}" alt="{target.name}"></div>'
            '<script>document.addEventListener("keydown",e=>{'
            'if(e.key==="ArrowLeft"){const a=document.getElementById("prev");if(a)location=a.href}'
            'if(e.key==="ArrowRight"){const a=document.getElementById("next");if(a)location=a.href}'
            '});</script>' + VIEWER_JS)
        return web.Response(text=html, content_type='text/html')

    return web.FileResponse(target, headers={'Cache-Control': 'no-cache'})


def main():
    # client_max_size: limit żądania POST (upload) — domyślny 1 MB to za mało
    app = web.Application(middlewares=[auth], client_max_size=512 * 1024 ** 2)
    app.on_startup.append(_lektor_restore)
    app.router.add_get('/{path:.*}', serve)
    app.router.add_post('/{path:.*}', upload)
    app.router.add_delete('/{path:.*}', delete_item)
    auth_info = f'user={USER}' if PASSW else 'OPEN (no auth)'
    print(f'sprawozdania-server: http://{HOST}:{PORT}/ — {auth_info}', flush=True)
    web.run_app(app, host=HOST, port=PORT, access_log=None)


if __name__ == '__main__':
    main()
